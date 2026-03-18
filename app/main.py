# ══════════════════════════════════════════════════════════════════════════════
# script.py  ·  ANE Análisis de Comentarios  v2.1
# ══════════════════════════════════════════════════════════════════════════════
import copy
import streamlit as st
import pandas as pd
import time
import os
import json
from io import BytesIO
from collections import Counter
from dotenv import load_dotenv
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types

# ── Inicialización ─────────────────────────────────────────────────────────────
load_dotenv()
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

st.set_page_config(
    page_title="ANE — Análisis de Comentarios",
    layout="wide",
    page_icon="📊",
    initial_sidebar_state="collapsed",
)

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTES
# ══════════════════════════════════════════════════════════════════════════════
TOPIC_PALETTE = [
    "#4f6ef7", "#22c55e", "#f97316", "#a855f7", "#06b6d4",
    "#eab308", "#ef4444", "#ec4899", "#14b8a6", "#84cc16",
]

POSTURA_COLORS = {
    "Soporte":              "#22c55e",
    "Rechazo / Objeción":   "#ef4444",
    "Propuesta de ajuste":  "#eab308",
    "Técnico / Neutro":     "#64748b",
}

_CHART_BASE = dict(
    paper_bgcolor="rgba(0,0,0,0)",
    plot_bgcolor="rgba(0,0,0,0)",
    font=dict(color="#e2e8f0", family="Arial, sans-serif"),
    margin=dict(t=56, b=36, l=36, r=36),
)
_AXIS_STYLE = dict(gridcolor="#2e3147", zerolinecolor="#2e3147", linecolor="#2e3147")


def _chart_layout(**extra):
    return {**_CHART_BASE, **extra}


# ══════════════════════════════════════════════════════════════════════════════
# CSS GLOBAL
# ══════════════════════════════════════════════════════════════════════════════
def inject_css():
    st.markdown("""
    <style>
    .stApp { background-color: #0f1117; }
    section[data-testid="stSidebar"] { background-color: #1a1d27; }
    header[data-testid="stHeader"] { background: transparent; }
    .block-container { padding-top: 2rem; }
    hr { border-color: #2e3147 !important; margin: 2rem 0; }
    .panorama-sub { color: #94a3b8; font-size: 14px; margin: -4px 0 24px; line-height: 1.6; }
    [data-testid="stExpander"] {
        border: none !important;
        background: #1a1d27 !important;
        border-radius: 0 0 10px 10px !important;
        margin-top: 0 !important;
    }
    [data-testid="stExpander"] summary {
        color: #64748b !important;
        font-size: 12px !important;
        padding: 8px 16px !important;
        background: #1a1d27 !important;
    }
    @keyframes pulse-tip {
        0%   { box-shadow: 0 0 0 0    rgba(167,139,250,.85); transform: translateY(-50%) scale(1);   }
        55%  { box-shadow: 0 0 0 10px rgba(167,139,250,0);   transform: translateY(-50%) scale(1.25);}
        100% { box-shadow: 0 0 0 0    rgba(167,139,250,0);   transform: translateY(-50%) scale(1);   }
    }
    .ane-tip {
        position: absolute;
        right: -7px; top: 50%;
        transform: translateY(-50%);
        width: 16px; height: 16px;
        background: #a78bfa;
        border-radius: 50%;
        animation: pulse-tip 1.5s ease-out infinite;
    }
    </style>
    """, unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# BARRA DE PROGRESO ANIMADA
# ══════════════════════════════════════════════════════════════════════════════
def _build_progress_html(pct: int, label: str, step: int, total_steps: int) -> str:
    return f"""
    <div style="
        background:#1a1d27; border:1px solid #2e3147; border-radius:12px;
        padding:18px 24px 20px; margin:12px 0; font-family:Arial,sans-serif;
    ">
        <div style="display:flex; justify-content:space-between; align-items:center; margin-bottom:12px;">
            <span style="color:#e2e8f0; font-size:14px; font-weight:600;">{label}</span>
            <span style="color:#a78bfa; font-size:13px; font-weight:700;">{pct}%</span>
        </div>
        <div style="background:#2e3147; border-radius:99px; height:8px; position:relative; overflow:visible;">
            <div style="
                background:linear-gradient(90deg,#4f6ef7,#7c3aed);
                border-radius:99px; height:8px;
                width:{pct}%; position:relative;
            ">
                <div class="ane-tip"></div>
            </div>
        </div>
        <div style="margin-top:10px; color:#64748b; font-size:12px;">
            Paso {step} de {total_steps}
        </div>
    </div>
    """


def render_progress(placeholder, pct: int, label: str, step: int, total: int):
    placeholder.markdown(_build_progress_html(pct, label, step, total), unsafe_allow_html=True)


def smooth_advance(placeholder, from_pct: int, to_pct: int,
                   label: str, step: int, total: int, tick: float = 0.06):
    delta = to_pct - from_pct
    if delta <= 0:
        render_progress(placeholder, to_pct, label, step, total)
        return
    step_size = max(1, delta // min(delta, 14))
    current = from_pct
    while current < to_pct:
        current = min(current + step_size, to_pct)
        render_progress(placeholder, current, label, step, total)
        time.sleep(tick)


# ══════════════════════════════════════════════════════════════════════════════
# FUNCIONES DE APOYO — Excel
# ══════════════════════════════════════════════════════════════════════════════
def validate_excel(df: pd.DataFrame) -> bool:
    options = [
        ['No.', 'Remitente', 'Observacion recibida'],
        ['No.', 'Remitente', 'Observación recibida'],
    ]
    return any(all(c in df.columns for c in opt) for opt in options)


def get_obs_col(df: pd.DataFrame):
    for name in ['Observación recibida', 'Observacion recibida']:
        if name in df.columns:
            return name
    return None


def read_excel_smart(uploaded_file) -> pd.DataFrame | None:
    for h in [23, 0, 1, 2]:
        try:
            uploaded_file.seek(0)
            df = pd.read_excel(uploaded_file, header=h)
            df.columns = [str(c).strip() for c in df.columns]
            if validate_excel(df):
                return df
        except Exception:
            continue
    return None


def _safe_int(val) -> int | None:
    """Convierte a int de forma segura: maneja float, str y NaN."""
    try:
        return int(float(val))
    except (ValueError, TypeError):
        return None


# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICAS — export amigable para impresión (fondo blanco, texto negro)
# ══════════════════════════════════════════════════════════════════════════════
def make_print_friendly(fig) -> go.Figure:
    """
    Clona la figura y aplica layout legible sobre fondo blanco para el DOCX.
    No modifica la figura original que se muestra en la UI.
    """
    fig2 = copy.deepcopy(fig)
    fig2.update_layout(
        paper_bgcolor="white",
        plot_bgcolor="#f0f2f5",
        font=dict(color="#1a1a1a", family="Arial, sans-serif", size=12),
        legend=dict(font=dict(color="#1a1a1a", size=11)),
        title_font=dict(color="#1a1a1a"),
    )
    fig2.update_xaxes(
        gridcolor="#cccccc",
        linecolor="#aaaaaa",
        zerolinecolor="#aaaaaa",
        tickfont=dict(color="#1a1a1a"),
        title_font=dict(color="#1a1a1a"),
    )
    fig2.update_yaxes(
        gridcolor="#cccccc",
        linecolor="#aaaaaa",
        zerolinecolor="#aaaaaa",
        tickfont=dict(color="#1a1a1a"),
        title_font=dict(color="#1a1a1a"),
    )
    # Corregir textfont de trazas (barras con etiquetas)
    for trace in fig2.data:
        if hasattr(trace, 'textfont') and trace.textfont is not None:
            trace.update(textfont=dict(color="#1a1a1a"))
    return fig2


def fig_to_bytes(fig, print_mode: bool = False) -> bytes:
    """
    Convierte figura Plotly a PNG.
    print_mode=True aplica fondo blanco y texto negro (para DOCX).
    """
    import plotly.io as pio
    f = make_print_friendly(fig) if print_mode else fig
    return pio.to_image(f, format="png", width=960, height=520, scale=2)


# ══════════════════════════════════════════════════════════════════════════════
# FUNCIONES IA
# ══════════════════════════════════════════════════════════════════════════════
def limpiar_json(texto: str) -> str:
    texto = texto.strip()
    if "```" in texto:
        for p in texto.split("```"):
            p2 = p.strip().lstrip("json").strip()
            try:
                json.loads(p2)
                return p2
            except Exception:
                continue
    return texto


def llamar_gemini(prompt_text: str) -> str:
    """
    Sin límite de max_output_tokens — igual que el script original.
    temperature=0.1 para respuestas deterministas y consistentes.
    """
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt_text,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            temperature=0.1,
        ),
    )
    return response.text


def procesar_con_ia(df: pd.DataFrame, obs_col: str):
    """
    Prompts 1 y 2: identificación de temas y specs de gráficas.
    Lógica sin cambios respecto a la versión original.
    """
    df_clean = df.dropna(subset=[obs_col])
    comentarios_muestra = (
        df_clean[['No.', 'Remitente', obs_col]]
        .rename(columns={obs_col: 'Observacion recibida'})
        .to_dict(orient='records')
    )

    # ── Prompt 1: Temas ────────────────────────────────────────────────────
    # Instrucción estricta de cantidad para que temperature=0.1 sea efectiva
    prompt_temas = f"""
Actua como un analista experto. Lee los siguientes comentarios de una consulta publica.
Tu tarea:
1. Identifica EXACTAMENTE entre 6 y 8 temas principales. No menos de 6, no mas de 8.
   Elige el numero que mejor represente la diversidad real del contenido.
2. Resume cada tema en texto plano SIN formato markdown (sin asteriscos, numerales ni guiones).
3. Asigna los IDs ("No.") de los comentarios a cada tema. Un comentario puede pertenecer a varios temas.

Devuelve SOLO un JSON con esta estructura exacta, sin texto adicional antes ni despues:
{{
  "temas": [
    {{
      "nombre": "Nombre del Tema",
      "resumen": "Resumen en texto plano...",
      "ids": [1, 2, 3]
    }}
  ]
}}

Comentarios:
{json.dumps(comentarios_muestra[:100], ensure_ascii=False)}
"""
    temas_data = json.loads(limpiar_json(llamar_gemini(prompt_temas)))['temas']
    time.sleep(15)

    # ── Prompt 2: Specs de gráficas ────────────────────────────────────────
    resumen_datos = {
        "total_comentarios": len(df_clean),
        "total_remitentes_unicos": df_clean['Remitente'].nunique(),
        "top_remitentes": df_clean['Remitente'].value_counts().head(5).to_dict(),
        "total_temas": len(temas_data),
        "temas": [{"nombre": t["nombre"], "cantidad": len(t["ids"])} for t in temas_data],
    }

    prompt_graficas = f"""
Eres un analista de datos. Resumen del dataset:
{json.dumps(resumen_datos, ensure_ascii=False)}

Fuentes disponibles: "temas", "remitente", "longitud_comentarios", "comentarios_por_tema_remitente"
Tipos disponibles: "pie", "bar", "treemap", "histogram", "box", "imshow"

Elige las 4 graficas mas informativas. Escribe un parrafo de analisis en texto plano para cada una.

Devuelve SOLO JSON:
{{
  "graficas": [
    {{
      "titulo": "...",
      "tipo": "...",
      "fuente_datos": "...",
      "top_n": 10,
      "analisis": "..."
    }}
  ]
}}
"""
    time.sleep(5)
    specs = json.loads(limpiar_json(llamar_gemini(prompt_graficas)))['graficas']
    return temas_data, specs


def procesar_posturas(df: pd.DataFrame, obs_col: str) -> dict:
    """
    Prompt 3: clasifica la postura de cada comentario.
    Agnóstico al dominio: funciona para cualquier temática.
    """
    df_clean = df.dropna(subset=[obs_col])
    comentarios = [
        {"id": _safe_int(row['No.']), "texto": str(row[obs_col])[:500]}
        for _, row in df_clean.iterrows()
        if _safe_int(row['No.']) is not None
    ]

    prompt = f"""
Clasifica la POSTURA de cada comentario eligiendo EXACTAMENTE una de estas 4 opciones:
"Soporte", "Rechazo / Objeción", "Propuesta de ajuste", "Técnico / Neutro"

Criterios:
- Soporte: apoya o valora positivamente la propuesta o medida
- Rechazo / Objeción: rechaza, critica o se opone directamente
- Propuesta de ajuste: propone cambios concretos sin rechazar del todo
- Técnico / Neutro: preguntas, aclaraciones o datos sin posición clara

Devuelve SOLO este JSON sin texto adicional:
{{"posturas": [{{"id": 1, "postura": "Soporte"}}]}}

Comentarios:
{json.dumps(comentarios, ensure_ascii=False)}
"""
    try:
        datos = json.loads(limpiar_json(llamar_gemini(prompt)))
        posturas_dict = {int(item['id']): item['postura'] for item in datos.get('posturas', [])}
    except Exception as e:
        st.warning(f"No se pudieron obtener las posturas automáticamente: {e}")
        posturas_dict = {}

    for _, row in df_clean.iterrows():
        key = _safe_int(row['No.'])
        if key is not None and key not in posturas_dict:
            posturas_dict[key] = 'Técnico / Neutro'

    return posturas_dict


def extraer_nombre_proyecto(temas_data: list) -> str:
    nombres = [t['nombre'] for t in temas_data[:5]]
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=(
            f"En máximo 12 palabras escribe el título del proyecto o temática que agrupa "
            f"estos temas: {nombres}. Responde SOLO el título, sin comillas ni punto final."
        ),
    )
    return response.text.strip()


# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICAS AI-DECIDIDAS  (para el DOCX — lógica original sin cambios)
# ══════════════════════════════════════════════════════════════════════════════
def construir_grafica(spec: dict, df: pd.DataFrame, temas_data: list, obs_col: str):
    tipo   = spec.get("tipo", "bar")
    titulo = spec.get("titulo", "Gráfica")
    fuente = spec.get("fuente_datos", "")
    fig = None
    try:
        if fuente == "temas":
            nombres    = [t['nombre'] for t in temas_data]
            cantidades = [len(t['ids']) for t in temas_data]
            df_plot = pd.DataFrame({"Tema": nombres, "Cantidad": cantidades})
            if tipo == "pie":
                fig = px.pie(df_plot, values="Cantidad", names="Tema", title=titulo, hole=0.4,
                             color_discrete_sequence=px.colors.sequential.Magma)
            elif tipo == "treemap":
                fig = px.treemap(df_plot, path=["Tema"], values="Cantidad", title=titulo)
            else:
                fig = px.bar(df_plot, x="Tema", y="Cantidad", title=titulo,
                             color="Cantidad", color_continuous_scale="Viridis")
                fig.update_layout(xaxis_tickangle=-30)

        elif fuente == "remitente":
            top_n = spec.get("top_n", 10)
            conteo = df['Remitente'].value_counts().head(top_n).reset_index()
            conteo.columns = ['Remitente', 'Cantidad']
            if tipo == "pie":
                fig = px.pie(conteo, values="Cantidad", names="Remitente", title=titulo)
            else:
                fig = px.bar(conteo, x="Cantidad", y="Remitente", orientation='h', title=titulo,
                             color="Cantidad", color_continuous_scale="Blues")
                fig.update_layout(yaxis={'categoryorder': 'total ascending'})

        elif fuente == "longitud_comentarios":
            df_len = df.copy()
            df_len['longitud'] = df_len[obs_col].astype(str).apply(len)
            if tipo == "box":
                fig = px.box(df_len, y="longitud", title=titulo)
            else:
                fig = px.histogram(df_len, x="longitud", title=titulo, nbins=30,
                                   color_discrete_sequence=["#636EFA"])

        elif fuente == "comentarios_por_tema_remitente":
            filas = []
            for t in temas_data:
                ids_t = {_safe_int(x) for x in t['ids']} - {None}
                df_t = df[df['No.'].apply(lambda v: _safe_int(v) in ids_t)]
                for rem in df_t['Remitente'].value_counts().head(5).index:
                    filas.append({"Tema": t['nombre'][:30], "Remitente": rem,
                                  "Cantidad": int(df_t[df_t['Remitente'] == rem].shape[0])})
            if filas:
                pivot = pd.DataFrame(filas).pivot_table(
                    index="Remitente", columns="Tema", values="Cantidad", fill_value=0)
                fig = px.imshow(pivot, title=titulo, color_continuous_scale="YlOrRd", aspect="auto")

    except Exception as e:
        st.warning(f"No se pudo generar '{titulo}': {e}")

    if fig:
        fig.update_layout(margin=dict(t=60, b=40, l=40, r=40))
    return fig


# ══════════════════════════════════════════════════════════════════════════════
# GRÁFICAS FIJAS  (interfaz — no dependen de decisiones IA)
# ══════════════════════════════════════════════════════════════════════════════
def crear_graficas_fijas(df: pd.DataFrame, obs_col: str,
                         temas_data: list, posturas_dict: dict) -> dict:
    charts = {}

    # ── 1. Donut: distribución de posturas ────────────────────────────────
    if posturas_dict:
        counts = Counter(posturas_dict.values())
        labels = list(counts.keys())
        values = list(counts.values())
        colors = [POSTURA_COLORS.get(l, '#888') for l in labels]
        fig1 = go.Figure(go.Pie(
            labels=labels, values=values, hole=0.58,
            marker_colors=colors,
            textfont=dict(size=13),
        ))
        fig1.update_layout(
            title="Distribución de Posturas",
            legend=dict(orientation="v", x=1.02, y=0.5, font=dict(size=12)),
            **_chart_layout(),
        )
        charts['donut_posturas'] = fig1

    # ── 2. Bar horizontal: comentarios por tema ───────────────────────────
    nombres    = [t['nombre'] for t in temas_data]
    cantidades = [len(t['ids']) for t in temas_data]
    colors_t   = [TOPIC_PALETTE[i % len(TOPIC_PALETTE)] for i in range(len(temas_data))]

    fig2 = go.Figure(go.Bar(
        x=cantidades, y=[n[:38] for n in nombres],
        orientation='h',
        marker_color=colors_t,
        text=cantidades, textposition='outside',
        textfont=dict(color='#e2e8f0'),
    ))
    fig2.update_layout(
        title="Comentarios por Tema",
        xaxis=dict(title="Cantidad", **_AXIS_STYLE),
        yaxis=dict(categoryorder='total ascending', **_AXIS_STYLE),
        **_chart_layout(),
    )
    charts['bar_temas'] = fig2

    # ── 3. Bar horizontal: top 10 remitentes ──────────────────────────────
    top_rem = df['Remitente'].value_counts().head(10).reset_index()
    top_rem.columns = ['Remitente', 'Cantidad']
    top_rem = top_rem.sort_values('Cantidad', ascending=True)

    fig3 = go.Figure(go.Bar(
        x=top_rem['Cantidad'], y=top_rem['Remitente'],
        orientation='h',
        marker=dict(color=top_rem['Cantidad'], colorscale='Blues', showscale=False),
        text=top_rem['Cantidad'], textposition='outside',
        textfont=dict(color='#e2e8f0'),
    ))
    fig3.update_layout(
        title="Top 10 Remitentes más Activos",
        xaxis=dict(**_AXIS_STYLE),
        yaxis=dict(categoryorder='total ascending', **_AXIS_STYLE),
        **_chart_layout(),
    )
    charts['bar_remitentes'] = fig3

    # ── 4. Box plot: longitud de comentarios por tema ──────────────────────
    rows_box = []
    for i, t in enumerate(temas_data):
        ids_t = {_safe_int(x) for x in t['ids']} - {None}
        df_t  = df[df['No.'].apply(lambda v: _safe_int(v) in ids_t)]
        for _, row in df_t.iterrows():
            rows_box.append({
                'Tema':     t['nombre'][:30],
                'Longitud': len(str(row[obs_col])),
            })
    if rows_box:
        df_box = pd.DataFrame(rows_box)
        fig4 = px.box(
            df_box, x='Tema', y='Longitud',
            title='Profundidad de Comentarios por Tema (nº de caracteres)',
            color='Tema',
            color_discrete_sequence=TOPIC_PALETTE,
            points='outliers',
        )
        fig4.update_layout(
            xaxis=dict(tickangle=-22, **_AXIS_STYLE),
            yaxis=dict(title="Caracteres", **_AXIS_STYLE),
            showlegend=False,
            **_chart_layout(),
        )
        charts['box_longitud'] = fig4

    # ── 5. Stacked bar: posturas por tema ─────────────────────────────────
    if posturas_dict:
        rows_p = []
        for t in temas_data:
            for id_ in t['ids']:
                key = _safe_int(id_)
                if key is not None:
                    rows_p.append({
                        'Tema':    t['nombre'][:30],
                        'Postura': posturas_dict.get(key, 'Técnico / Neutro'),
                    })
        pivot_p = (
            pd.DataFrame(rows_p)
            .groupby(['Tema', 'Postura']).size()
            .reset_index(name='N')
        )
        fig5 = px.bar(
            pivot_p, x='Tema', y='N', color='Postura',
            title='Distribución de Posturas por Tema',
            barmode='stack',
            color_discrete_map=POSTURA_COLORS,
        )
        fig5.update_layout(
            xaxis=dict(tickangle=-22, **_AXIS_STYLE),
            yaxis=dict(**_AXIS_STYLE),
            legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
            **_chart_layout(),
        )
        charts['stacked_posturas'] = fig5

    # ── 6. Heatmap: participación remitentes × temas ──────────────────────
    top_rems = df['Remitente'].value_counts().head(10).index.tolist()
    filas_h = []
    for t in temas_data:
        ids_t = {_safe_int(x) for x in t['ids']} - {None}
        df_t  = df[df['No.'].apply(lambda v: _safe_int(v) in ids_t)]
        for rem in top_rems:
            cnt = int(df_t[df_t['Remitente'] == rem].shape[0])
            if cnt > 0:
                filas_h.append({'Remitente': rem, 'Tema': t['nombre'][:25], 'N': cnt})
    if filas_h:
        pivot_h = (
            pd.DataFrame(filas_h)
            .pivot_table(index='Remitente', columns='Tema', values='N', fill_value=0)
        )
        fig6 = px.imshow(
            pivot_h,
            title='Participación de Remitentes por Tema',
            color_continuous_scale='Blues',
            aspect='auto',
            text_auto=True,
        )
        fig6.update_layout(
            xaxis=dict(tickangle=-22, gridcolor='#2e3147'),
            yaxis=dict(gridcolor='#2e3147'),
            **_chart_layout(),
        )
        charts['heatmap'] = fig6

    return charts


# ══════════════════════════════════════════════════════════════════════════════
# TABLA DE REMITENTES ENRIQUECIDA
# ══════════════════════════════════════════════════════════════════════════════
def crear_tabla_remitentes(df: pd.DataFrame, temas_data: list,
                           posturas_dict: dict) -> pd.DataFrame:
    """
    FIX: normaliza todos los IDs con _safe_int() para evitar
    que float vs int rompa la intersección de conjuntos.
    """
    # Precompute: id → lista de nombres de tema
    id_to_temas: dict[int, list[str]] = {}
    for t in temas_data:
        for raw_id in t['ids']:
            key = _safe_int(raw_id)
            if key is not None:
                id_to_temas.setdefault(key, []).append(t['nombre'])

    filas = []
    for remitente in df['Remitente'].dropna().unique():
        sub     = df[df['Remitente'] == remitente]
        ids_int = [k for k in (
            _safe_int(v) for v in sub['No.'].tolist()
        ) if k is not None]

        post_list   = [posturas_dict.get(i, 'Técnico / Neutro') for i in ids_int]
        postura_dom = Counter(post_list).most_common(1)[0][0] if post_list else 'Técnico / Neutro'

        tema_count: Counter = Counter()
        for id_ in ids_int:
            for tname in id_to_temas.get(id_, []):
                tema_count[tname] += 1

        top_temas = ' · '.join(n[:24] for n, _ in tema_count.most_common(2)) or '—'

        filas.append({
            'Remitente':         remitente,
            'Comentarios':       len(ids_int),
            'Postura Dominante': postura_dom,
            'Temas Principales': top_temas,
        })

    return (
        pd.DataFrame(filas)
        .sort_values('Comentarios', ascending=False)
        .reset_index(drop=True)
    )


# ══════════════════════════════════════════════════════════════════════════════
# GENERACIÓN DOCX  — fuente Arial 12 negra en todos los elementos
# ══════════════════════════════════════════════════════════════════════════════
def _set_run_black(run, size: int = 11, bold: bool = False):
    """Aplica Arial, tamaño, negro y negrita a un run de python-docx."""
    run.font.name  = 'Arial'
    run.font.size  = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold


def _add_doc_heading(doc, text: str, size: int = 13) -> None:
    """
    Añade un párrafo de encabezado con Arial negro en lugar del estilo
    de Word que hereda el color azul del tema del documento.
    """
    p   = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_black(run, size=size, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)


def _add_bullet_safe(doc, label: str, body: str) -> None:
    try:
        p = doc.add_paragraph(style='List Bullet')
    except KeyError:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
    r_label = p.add_run(f"{label} ")
    _set_run_black(r_label, bold=True)
    r_body = p.add_run(body)
    _set_run_black(r_body, bold=False)


def generar_docx(temas_data: list, graficas_info: list, titulo_informe: str,
                 posturas_dict: dict | None = None, nombre_proyecto: str = "") -> bytes:
    doc = Document()

    # Estilo base del documento
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Portada ─────────────────────────────────────────────────────────────
    # Título principal en negro (evita herencia de color azul del tema Word)
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_titulo = p_titulo.add_run(titulo_informe)
    _set_run_black(r_titulo, size=18, bold=True)

    if nombre_proyecto:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_sub = p_sub.add_run(nombre_proyecto)
        _set_run_black(r_sub, size=12, bold=False)
        r_sub.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # ── Sección 1: Posturas ─────────────────────────────────────────────────
    if posturas_dict:
        _add_doc_heading(doc, '1. Distribución de Posturas', size=13)
        counts = Counter(posturas_dict.values())
        total  = sum(counts.values())
        for postura, cnt in sorted(counts.items(), key=lambda x: -x[1]):
            _add_bullet_safe(doc, f'{postura}:',
                             f'{cnt} comentarios ({cnt / total * 100:.1f}%)')
        doc.add_paragraph()

    # ── Sección 2: Temas ────────────────────────────────────────────────────
    _add_doc_heading(doc, '2. Temas Principales Identificados', size=13)
    for tema in temas_data:
        p_tema = doc.add_paragraph()
        r_tema = p_tema.add_run(f"{tema['nombre']}  ({len(tema['ids'])} comentarios)")
        _set_run_black(r_tema, size=11, bold=True)
        p_tema.paragraph_format.space_before = Pt(10)

        p_body = doc.add_paragraph(tema['resumen'].replace('**', '').replace('#', ''))
        p_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p_body.runs:
            _set_run_black(run, size=11, bold=False)

    doc.add_paragraph()

    # ── Sección 3: Gráficas ─────────────────────────────────────────────────
    _add_doc_heading(doc, '3. Análisis Cuantitativo y Visualizaciones', size=13)
    for g in graficas_info:
        p_gtitle = doc.add_paragraph()
        r_gt = p_gtitle.add_run(g['titulo'])
        _set_run_black(r_gt, size=11, bold=True)
        p_gtitle.paragraph_format.space_before = Pt(10)

        if g.get('analisis'):
            p_anal = doc.add_paragraph(g['analisis'].replace('**', '').replace('#', ''))
            p_anal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p_anal.runs:
                _set_run_black(run, size=11, bold=False)

        if g.get('imagen_bytes'):
            try:
                doc.add_picture(BytesIO(g['imagen_bytes']), width=Inches(5.8))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                p_err = doc.add_paragraph(f"[Gráfica no disponible: {e}]")
                for run in p_err.runs:
                    _set_run_black(run)

        doc.add_paragraph()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS DE INTERFAZ
# ══════════════════════════════════════════════════════════════════════════════
def render_kpis(n_comentarios: int, n_remitentes: int,
                n_temas: int, nombre_proyecto: str):
    st.markdown(f"### {nombre_proyecto}")
    st.markdown(
        f"<p class='panorama-sub'>"
        f"<b style='color:#e2e8f0'>{n_comentarios}</b> comentarios formales de "
        f"<b style='color:#e2e8f0'>{n_remitentes}</b> remitentes sobre el proyecto"
        f"</p>",
        unsafe_allow_html=True,
    )

    def kpi_card(val, lbl):
        return (
            f"<div style='background:#1a1d27;border:1px solid #2e3147;border-radius:12px;"
            f"padding:20px 24px;text-align:center;'>"
            f"<div style='font-size:38px;font-weight:800;color:#4f6ef7;line-height:1.2'>{val}</div>"
            f"<div style='font-size:11px;color:#94a3b8;margin-top:6px;"
            f"text-transform:uppercase;letter-spacing:.6px'>{lbl}</div>"
            f"</div>"
        )

    c1, c2, c3 = st.columns(3)
    c1.markdown(kpi_card(n_comentarios, "Comentarios analizados"), unsafe_allow_html=True)
    c2.markdown(kpi_card(n_remitentes,  "Remitentes únicos"),      unsafe_allow_html=True)
    c3.markdown(kpi_card(n_temas,       "Temas identificados"),     unsafe_allow_html=True)


def render_topic_cards(temas_data: list, df: pd.DataFrame,
                       obs_col: str, posturas_dict: dict):
    st.header("Temas Identificados")
    for i, tema in enumerate(temas_data):
        color = TOPIC_PALETTE[i % len(TOPIC_PALETTE)]

        st.markdown(f"""
        <div style="
            border-left:4px solid {color};
            background:{color}0d;
            border-radius:0 10px 10px 0;
            padding:14px 20px; margin-bottom:2px;
        ">
            <div style="display:flex;justify-content:space-between;
                        align-items:flex-start;flex-wrap:wrap;gap:8px;">
                <span style="font-size:15px;font-weight:700;color:#e2e8f0;">
                    {tema['nombre']}
                </span>
                <span style="background:{color}33;color:{color};
                             padding:3px 12px;border-radius:20px;
                             font-size:12px;font-weight:600;white-space:nowrap;">
                    {len(tema['ids'])} comentarios
                </span>
            </div>
            <p style="color:#94a3b8;font-size:13px;margin:8px 0 0 0;line-height:1.65;">
                {tema['resumen'][:360]}{'…' if len(tema['resumen']) > 360 else ''}
            </p>
        </div>
        """, unsafe_allow_html=True)

        with st.expander(f"Ver comentarios — {tema['nombre'][:45]}"):
            ids_t  = {_safe_int(x) for x in tema['ids']} - {None}
            df_t   = df[df['No.'].apply(lambda v: _safe_int(v) in ids_t)].head(3)
            for _, row in df_t.iterrows():
                postura = posturas_dict.get(_safe_int(row['No.']), 'Técnico / Neutro')
                p_color = POSTURA_COLORS.get(postura, '#64748b')
                texto   = str(row[obs_col])
                st.markdown(f"""
                <div style="padding:12px 14px;border-radius:8px;
                            background:rgba(255,255,255,0.04);
                            border:1px solid rgba(255,255,255,0.06);
                            margin-bottom:8px;">
                    <div style="display:flex;justify-content:space-between;
                                margin-bottom:5px;flex-wrap:wrap;gap:4px;">
                        <strong style="color:#e2e8f0">{row['Remitente']}</strong>
                        <span style="color:{p_color};font-size:12px;font-weight:600">
                            {postura}
                        </span>
                    </div>
                    <p style="color:#94a3b8;font-size:13px;margin:0;line-height:1.65">
                        {texto[:560]}{'…' if len(texto) > 560 else ''}
                    </p>
                </div>
                """, unsafe_allow_html=True)

        st.markdown("<div style='height:6px'></div>", unsafe_allow_html=True)


def render_charts(charts: dict):
    st.header("Análisis Visual")

    c_left, c_right = st.columns([1, 1.5])
    if 'donut_posturas' in charts:
        c_left.plotly_chart(charts['donut_posturas'], use_container_width=True)
    if 'bar_temas' in charts:
        c_right.plotly_chart(charts['bar_temas'], use_container_width=True)

    c1, c2 = st.columns(2)
    if 'bar_remitentes' in charts:
        c1.plotly_chart(charts['bar_remitentes'], use_container_width=True)
    if 'box_longitud' in charts:
        c2.plotly_chart(charts['box_longitud'], use_container_width=True)

    if 'stacked_posturas' in charts:
        st.plotly_chart(charts['stacked_posturas'], use_container_width=True)

    if 'heatmap' in charts:
        st.plotly_chart(charts['heatmap'], use_container_width=True)


def render_remitente_table(df_tabla: pd.DataFrame):
    st.header("Tabla de Remitentes")
    st.dataframe(
        df_tabla,
        use_container_width=True,
        hide_index=True,
        height=min(620, 52 + len(df_tabla) * 40),
        column_config={
            "Remitente":         st.column_config.TextColumn("Remitente",          width="medium"),
            "Comentarios":       st.column_config.NumberColumn("Comentarios",       format="%d",   width="small"),
            "Postura Dominante": st.column_config.TextColumn("Postura Dominante",  width="medium"),
            "Temas Principales": st.column_config.TextColumn("Temas Principales",  width="large"),
        },
    )


# ══════════════════════════════════════════════════════════════════════════════
# APP PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════
def main():
    inject_css()

    st.title("ANE — Análisis de Comentarios")
    st.markdown(
        "Sube el Excel con los comentarios de consulta pública. "
        "La IA identificará temas, clasificará posturas y generará un informe Word."
    )

    nombre_input = st.text_input(
        "Nombre del proyecto  (opcional — si lo dejas vacío se genera automáticamente)",
        placeholder="Ej: Analisis de comentarios de ...",
    )

    uploaded_file = st.file_uploader(
        "Arrastra tu Excel aquí o búscalo en tu equipo",
        type=['xlsx'],
    )

    if uploaded_file is None:
        return

    file_sig = f"{uploaded_file.name}_{uploaded_file.size}"
    if st.session_state.get('_file_sig') != file_sig:
        for k in ['procesado', 'df', 'obs_col', 'temas_data', 'posturas_dict',
                  'charts_fijos', 'graficas_info', 'df_tabla_rem', 'docx_bytes',
                  'nombre_proyecto']:
            st.session_state.pop(k, None)
        st.session_state['_file_sig'] = file_sig

    if 'procesado' in st.session_state and nombre_input.strip():
        st.session_state['nombre_proyecto'] = nombre_input.strip()

    if 'procesado' not in st.session_state:

        df = read_excel_smart(uploaded_file)
        if df is None:
            st.error(
                "El archivo no tiene la estructura esperada. "
                "Se necesitan las columnas: No., Remitente, Observacion recibida."
            )
            return
        obs_col = get_obs_col(df)

        ph  = st.empty()
        TOT = 6

        smooth_advance(ph, 0, 10, "Estructurando datos del archivo…", 1, TOT)

        render_progress(ph, 10, "Identificando temas con IA  (puede tardar ~5 mins)…", 2, TOT)
        temas_data, specs_graficas = procesar_con_ia(df, obs_col)
        smooth_advance(ph, 10, 45, "Temas identificados ✓", 2, TOT)

        render_progress(ph, 45, "Analizando posturas de los comentarios…", 3, TOT)
        posturas_dict = procesar_posturas(df, obs_col)
        smooth_advance(ph, 45, 65, "Posturas analizadas ✓", 3, TOT)

        render_progress(ph, 65, "Determinando nombre del proyecto…", 4, TOT)
        nombre_proyecto = (
            nombre_input.strip() if nombre_input.strip()
            else extraer_nombre_proyecto(temas_data)
        )
        smooth_advance(ph, 65, 72, "Nombre del proyecto listo ✓", 4, TOT)

        smooth_advance(ph, 72, 78, "Construyendo gráficas dinámicas…", 5, TOT)

        # Gráficas AI para el DOCX
        graficas_info = []
        for spec in specs_graficas:
            fig = construir_grafica(spec, df, temas_data, obs_col)
            if fig:
                try:
                    img_bytes = fig_to_bytes(fig, print_mode=True)   # ← fondo blanco
                except Exception:
                    img_bytes = None
                graficas_info.append({
                    "titulo":       spec["titulo"],
                    "analisis":     spec["analisis"],
                    "fig":          fig,
                    "imagen_bytes": img_bytes,
                })

        # Gráficas fijas para la UI
        charts_fijos = crear_graficas_fijas(df, obs_col, temas_data, posturas_dict)

        # Tabla remitentes
        df_tabla_rem = crear_tabla_remitentes(df, temas_data, posturas_dict)

        smooth_advance(ph, 78, 88, "Gráficas construidas ✓", 5, TOT)

        smooth_advance(ph, 88, 92, "Generando informe Word…", 6, TOT)

        # Gráficas fijas clave para el DOCX (con print_mode=True)
        graficas_docx = [g for g in graficas_info if g.get("imagen_bytes")]
        for key, label, analisis in [
            ('donut_posturas',   'Distribución de Posturas',
             'Clasificación de la postura de todos los comentarios analizados.'),
            ('stacked_posturas', 'Posturas por Tema',
             'Desglose de posturas dentro de cada tema identificado.'),
            ('bar_temas',        'Comentarios por Tema',
             'Número de comentarios agrupados por cada tema.'),
            ('box_longitud',     'Profundidad de Comentarios por Tema',
             'Distribución de la longitud de los comentarios en cada tema.'),
        ]:
            if key in charts_fijos:
                try:
                    graficas_docx.append({
                        "titulo":       label,
                        "analisis":     analisis,
                        "imagen_bytes": fig_to_bytes(charts_fijos[key], print_mode=True),
                    })
                except Exception:
                    pass

        docx_bytes = generar_docx(
            temas_data, graficas_docx,
            "Informe de Análisis de Comentarios",
            posturas_dict=posturas_dict,
            nombre_proyecto=nombre_proyecto,
        )

        smooth_advance(ph, 92, 100, "¡Análisis completado!", 6, TOT)
        time.sleep(0.9)
        ph.empty()

        st.session_state.update({
            'procesado':       True,
            'df':              df,
            'obs_col':         obs_col,
            'temas_data':      temas_data,
            'posturas_dict':   posturas_dict,
            'charts_fijos':    charts_fijos,
            'graficas_info':   graficas_info,
            'df_tabla_rem':    df_tabla_rem,
            'docx_bytes':      docx_bytes,
            'nombre_proyecto': nombre_proyecto,
        })

    # ── Resultados ─────────────────────────────────────────────────────────
    ss = st.session_state
    st.success("Análisis completado con éxito.")

    render_kpis(
        n_comentarios=len(ss['df']),
        n_remitentes=ss['df']['Remitente'].nunique(),
        n_temas=len(ss['temas_data']),
        nombre_proyecto=ss['nombre_proyecto'],
    )
    st.divider()

    render_topic_cards(ss['temas_data'], ss['df'], ss['obs_col'], ss['posturas_dict'])
    st.divider()

    render_charts(ss['charts_fijos'])
    st.divider()

    render_remitente_table(ss['df_tabla_rem'])
    st.divider()

    st.download_button(
        label="⬇️ Descargar Informe Ejecutivo (DOCX)",
        data=ss['docx_bytes'],
        file_name="Informe_ANE_Comentarios.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="descarga_docx",
    )


if __name__ == "__main__":
    main()
import streamlit as st
import pandas as pd
import time
import os
import json
from io import BytesIO
from dotenv import load_dotenv
import plotly.express as px
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types

# ---------------- CONFIGURACION INICIAL ----------------
load_dotenv()
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

st.set_page_config(page_title="ANE Analisis de Comentarios", layout="wide")

# ---------------- FUNCIONES DE APOYO ----------------

def validate_excel(df):
    required_cols = ['No.', 'Remitente', 'Observacion recibida']
    # Intentar con tilde tambien
    alt_cols = ['No.', 'Remitente', 'Observación recibida']
    return (all(col in df.columns for col in required_cols) or
            all(col in df.columns for col in alt_cols))

def get_obs_col(df):
    """Detecta el nombre exacto de la columna de observaciones."""
    for name in ['Observación recibida', 'Observacion recibida']:
        if name in df.columns:
            return name
    return None

def fig_to_bytes(fig):
    """Convierte figura Plotly a PNG usando kaleido 0.2.1."""
    import plotly.io as pio
    img_bytes = pio.to_image(fig, format="png", width=800, height=450, scale=2)
    return img_bytes

def generar_docx(temas_data, graficas_info, titulo_informe):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    titulo = doc.add_heading(titulo_informe, 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Seccion 1: Temas
    h1 = doc.add_paragraph()
    r1 = h1.add_run('1. Resumen de Temas de Interes')
    r1.bold = True
    r1.font.size = Pt(14)

    for tema in temas_data:
        p = doc.add_paragraph()
        p.add_run(f"Tema: {tema['nombre']}").bold = True
        doc.add_paragraph(tema['resumen'].replace('**', '').replace('#', ''))

    doc.add_paragraph()

    # Seccion 2: Graficas
    h2 = doc.add_paragraph()
    r2 = h2.add_run('2. Analisis Cuantitativo y Graficas')
    r2.bold = True
    r2.font.size = Pt(14)

    for g in graficas_info:
        p = doc.add_paragraph()
        p.add_run(g['titulo']).bold = True
        doc.add_paragraph(g['analisis'].replace('**', '').replace('#', ''))
        try:
            img_bytes = BytesIO(g['imagen_bytes'])
            doc.add_picture(img_bytes, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"[No se pudo insertar la grafica: {e}]")
        doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def construir_grafica(spec, df, temas_data, obs_col):
    tipo = spec.get("tipo", "bar")
    titulo = spec.get("titulo", "Grafica")
    fuente = spec.get("fuente_datos", "")

    fig = None

    try:
        if fuente == "temas":
            nombres = [t['nombre'] for t in temas_data]
            cantidades = [len(t['ids']) for t in temas_data]
            df_plot = pd.DataFrame({"Tema": nombres, "Cantidad": cantidades})

            if tipo == "pie":
                fig = px.pie(df_plot, values="Cantidad", names="Tema",
                             title=titulo, hole=0.4,
                             color_discrete_sequence=px.colors.sequential.Magma)
            elif tipo == "treemap":
                fig = px.treemap(df_plot, path=["Tema"], values="Cantidad", title=titulo)
            else:  # bar por defecto
                fig = px.bar(df_plot, x="Tema", y="Cantidad", title=titulo,
                             color="Cantidad", color_continuous_scale="Viridis")
                fig.update_layout(xaxis_tickangle=-30)

        elif fuente == "remitente":
            top_n = spec.get("top_n", 10)
            conteo = df['Remitente'].value_counts().head(top_n).reset_index()
            conteo.columns = ['Remitente', 'Cantidad']

            if tipo == "pie":
                fig = px.pie(conteo, values="Cantidad", names="Remitente", title=titulo)
            else:
                fig = px.bar(conteo, x="Cantidad", y="Remitente", orientation='h',
                             title=titulo, color="Cantidad",
                             color_continuous_scale="Blues")
                fig.update_layout(yaxis={'categoryorder': 'total ascending'})

        elif fuente == "longitud_comentarios":
            df_len = df.copy()
            df_len['longitud'] = df_len[obs_col].astype(str).apply(len)
            if tipo == "box":
                fig = px.box(df_len, y="longitud", title=titulo)
            else:
                fig = px.histogram(df_len, x="longitud", title=titulo, nbins=30,
                                   color_discrete_sequence=["#636EFA"])

        elif fuente == "comentarios_por_tema_remitente":
            filas = []
            for t in temas_data:
                df_t = df[df['No.'].isin(t['ids'])]
                for rem in df_t['Remitente'].value_counts().head(5).index:
                    filas.append({
                        "Tema": t['nombre'][:30],
                        "Remitente": rem,
                        "Cantidad": int(df_t[df_t['Remitente'] == rem].shape[0])
                    })
            if filas:
                df_heat = pd.DataFrame(filas)
                pivot = df_heat.pivot_table(index="Remitente", columns="Tema",
                                            values="Cantidad", fill_value=0)
                fig = px.imshow(pivot, title=titulo, color_continuous_scale="YlOrRd",
                                aspect="auto")

    except Exception as e:
        st.warning(f"No se pudo generar la grafica '{titulo}': {e}")
        return None

    if fig:
        fig.update_layout(margin=dict(t=60, b=40, l=40, r=40))

    return fig

def limpiar_json(texto):
    texto = texto.strip()
    if "```" in texto:
        partes = texto.split("```")
        for p in partes:
            p2 = p.strip()
            if p2.startswith("json"):
                p2 = p2[4:].strip()
            try:
                json.loads(p2)
                return p2
            except Exception:
                continue
    return texto

def llamar_gemini(prompt_text):
    """Wrapper para la nueva API google.genai."""
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt_text,
        config=types.GenerateContentConfig(
            response_mime_type="application/json"
        )
    )
    return response.text

def procesar_con_ia(df, obs_col):
    df_clean = df.dropna(subset=[obs_col])
    comentarios_muestra = df_clean[['No.', 'Remitente', obs_col]].rename(
        columns={obs_col: 'Observacion recibida'}
    ).to_dict(orient='records')

    # ---- Prompt 1: Clasificacion tematica ----
    prompt_temas = f"""
Actua como un analista experto. Lee los siguientes comentarios de una consulta publica.
Tu tarea:
1. Agrupa los comentarios en 5 a 10 temas principales (los mas relevantes).
2. Resume cada tema en texto plano SIN usar formato markdown (sin asteriscos, sin numerales, sin guiones).
3. Asigna los IDs ("No.") de los comentarios a cada tema.

Devuelve SOLO un JSON con esta estructura exacta, sin texto adicional antes ni despues:
{{
  "temas": [
    {{
      "nombre": "Nombre del Tema",
      "resumen": "Resumen en texto plano sin formato...",
      "ids": [1, 2, 3]
    }}
  ]
}}

Comentarios:
{json.dumps(comentarios_muestra[:100], ensure_ascii=False)}
"""

    respuesta_temas = llamar_gemini(prompt_temas)
    datos_json = json.loads(limpiar_json(respuesta_temas))
    temas_data = datos_json['temas']

    time.sleep(15)

    # ---- Prompt 2: IA decide que graficas generar ----
    resumen_datos = {
        "total_comentarios": len(df_clean),
        "total_remitentes_unicos": df_clean['Remitente'].nunique(),
        "top_remitentes": df_clean['Remitente'].value_counts().head(5).to_dict(),
        "total_temas": len(temas_data),
        "temas": [{"nombre": t["nombre"], "cantidad": len(t["ids"])} for t in temas_data]
    }

    prompt_graficas = f"""
Eres un analista de datos. Tienes este resumen de un conjunto de comentarios de una consulta publica:
{json.dumps(resumen_datos, ensure_ascii=False)}

Las fuentes de datos disponibles para graficar son:
- "temas": distribucion de comentarios por tema
- "remitente": participacion de actores/entidades
- "longitud_comentarios": longitud en caracteres de los comentarios
- "comentarios_por_tema_remitente": cruce entre temas y remitentes (heatmap)

Los tipos de grafica disponibles son: "pie", "bar", "treemap", "histogram", "box", "imshow"

Decide cuales son las 4 graficas mas informativas para este conjunto de datos.
Para cada una escribe un parrafo de analisis en texto plano (sin markdown, sin asteriscos, sin viñetas).

Devuelve SOLO un JSON con esta estructura exacta, sin texto adicional:
{{
  "graficas": [
    {{
      "titulo": "Titulo descriptivo de la grafica",
      "tipo": "pie|bar|treemap|histogram|box|imshow",
      "fuente_datos": "temas|remitente|longitud_comentarios|comentarios_por_tema_remitente",
      "top_n": 10,
      "analisis": "Parrafo de analisis en texto plano..."
    }}
  ]
}}
"""

    time.sleep(5)
    respuesta_graficas = llamar_gemini(prompt_graficas)
    specs_graficas = json.loads(limpiar_json(respuesta_graficas))

    return temas_data, specs_graficas['graficas']


# ---------------- INTERFAZ DE USUARIO ----------------
st.title("ANE Analisis de Comentarios")
st.markdown(
    "Sube el archivo Excel con los comentarios de la consulta publica. "
    "La IA clasificara los temas, decidira que graficas generar y producira "
    "un informe Word con las visualizaciones incluidas."
)

uploaded_file = st.file_uploader(
    "Arrastra tu Excel aqui o buscalo en tu ordenador", type=['xlsx']
)

if uploaded_file is not None:
    if 'procesado' not in st.session_state:
        try:
            df = pd.read_excel(uploaded_file, header=23)
            df.columns = [str(c).strip() for c in df.columns]

            obs_col = get_obs_col(df)
            if obs_col is None or not validate_excel(df):
                st.error(
                    "El documento no tiene la estructura requerida. "
                    "Se esperan columnas: No., Remitente, Observacion recibida. "
                    "Verifica que los encabezados esten en la fila correcta."
                )
                st.stop()
        except Exception as e:
            st.error(f"Error al leer el Excel: {e}")
            st.stop()

        with st.spinner('Procesando con Inteligencia Artificial... Por favor espera (puede tomar cerca de 1 minuto).'):
            progress_bar = st.progress(0)

            progress_bar.progress(15, text="Estructurando datos...")
            temas_data, specs_graficas = procesar_con_ia(df, obs_col)

            progress_bar.progress(60, text="Construyendo graficas dinamicas...")
            graficas_info = []
            for spec in specs_graficas:
                fig = construir_grafica(spec, df, temas_data, obs_col)
                if fig:
                    try:
                        img_bytes = fig_to_bytes(fig)
                    except Exception as e:
                        st.warning(f"No se pudo exportar la grafica '{spec.get('titulo', '')}' como imagen: {e}")
                        img_bytes = None
                    graficas_info.append({
                        "titulo": spec["titulo"],
                        "analisis": spec["analisis"],
                        "fig": fig,
                        "imagen_bytes": img_bytes
                    })

            progress_bar.progress(85, text="Generando documento DOCX con graficas...")

            # Filtrar las que si tienen imagen para el docx
            graficas_para_docx = [g for g in graficas_info if g["imagen_bytes"] is not None]
            docx_bytes = generar_docx(temas_data, graficas_para_docx, "Informe de Analisis de Comentarios")

            progress_bar.progress(100, text="Proceso finalizado.")
            time.sleep(1)
            progress_bar.empty()

        st.session_state['procesado'] = True
        st.session_state['df'] = df
        st.session_state['obs_col'] = obs_col
        st.session_state['temas_data'] = temas_data
        st.session_state['graficas_info'] = graficas_info
        st.session_state['docx_bytes'] = docx_bytes

    # ---- Resultados ----
    st.success("Analisis completado con exito.")

    st.download_button(
        label="Descargar Informe Completo en DOCX",
        data=st.session_state['docx_bytes'],
        file_name="Informe_ANE_Comentarios.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="descarga_docx"
    )

    st.divider()

    # ---- Temas ----
    st.header("Temas de Interes Identificados")
    df_display = st.session_state['df']
    obs_col = st.session_state['obs_col']

    for tema in st.session_state['temas_data']:
        with st.expander(f"{tema['nombre']} ({len(tema['ids'])} comentarios)"):
            st.markdown(f"**Analisis:** {tema['resumen'].replace('**', '')}")
            st.markdown("---")
            st.markdown("**Comentarios asociados:**")
            df_tema = df_display[df_display['No.'].isin(tema['ids'])]
            for _, row in df_tema.iterrows():
                st.markdown(f"**{row['Remitente']}**: {row[obs_col]}")

    st.divider()

    # ---- Graficas ----
    st.header("Analitica Visual generada por IA")

    graficas_info = st.session_state['graficas_info']
    cols = st.columns(2)

    for i, g in enumerate(graficas_info):
        with cols[i % 2]:
            st.subheader(g['titulo'])
            st.plotly_chart(g['fig'], use_container_width=True)
            with st.popover("Ver interpretacion"):
                st.write(g['analisis'])